# """
# resume/parser.py
# ================
# Resume parsing for PDF and DOCX files.
# Uses Ollama (local Llama 3 / Mistral) to extract structured data — 100% offline.

# Flow:  uploaded file → extract raw text → Ollama LLM → structured dict
# """

# import json
# import logging
# from pathlib import Path
# from dotenv import load_dotenv
# from utils.ollama_client import chat
# from utils.prompts import get_resume_extraction_prompt

# load_dotenv()
# logger = logging.getLogger(__name__)


# # ── Text Extraction ──────────────────────────────────────────────────────────

# def _extract_pdf(file_path: str) -> str:
#     """Extract plain text from a PDF using PyPDF2."""
#     try:
#         import PyPDF2
#         parts = []
#         with open(file_path, "rb") as f:
#             reader = PyPDF2.PdfReader(f)
#             for i, page in enumerate(reader.pages):
#                 try:
#                     text = page.extract_text()
#                     if text:
#                         parts.append(text)
#                 except Exception as e:
#                     logger.warning(f"Skipping PDF page {i}: {e}")
#         return "\n".join(parts)
#     except Exception as e:
#         raise ValueError(f"PDF parsing failed: {e}")


# def _extract_docx(file_path: str) -> str:
#     """Extract plain text from a DOCX, including table cells."""
#     try:
#         from docx import Document
#         doc = Document(file_path)
#         parts = []

#         for para in doc.paragraphs:
#             if para.text.strip():
#                 parts.append(para.text.strip())

#         # Many resumes use tables for layout — include those too
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     if cell.text.strip():
#                         parts.append(cell.text.strip())

#         return "\n".join(parts)
#     except Exception as e:
#         raise ValueError(f"DOCX parsing failed: {e}")


# def _extract_text(file_path: str) -> str:
#     """Route to the right extractor based on file extension."""
#     ext = Path(file_path).suffix.lower()
#     if ext == ".pdf":
#         return _extract_pdf(file_path)
#     elif ext in (".docx", ".doc"):
#         return _extract_docx(file_path)
#     else:
#         raise ValueError(f"Unsupported file type '{ext}'. Please upload PDF or DOCX.")


# # ── AI Extraction ─────────────────────────────────────────────────────────────

# def _extract_with_ai(raw_text: str) -> dict:
#     """
#     Pass raw resume text to the local Ollama model and parse structured JSON.
#     Uses a very low temperature for deterministic extraction.
#     """
#     prompt = get_resume_extraction_prompt(raw_text)

#     content = chat(
#         messages=[{"role": "user", "content": prompt}],
#         system_prompt=(
#             "You are a precise resume parser. "
#             "Always respond with valid JSON only. "
#             "No markdown fences, no explanations, no extra text."
#         ),
#         temperature=0.1,   # Near-deterministic for data extraction
#         max_tokens=1500,
#     )

#     # Strip markdown fences if the model added them anyway
#     cleaned = content.strip()
#     if cleaned.startswith("```"):
#         lines = cleaned.split("\n")
#         cleaned = "\n".join(
#             line for line in lines
#             if not line.strip().startswith("```")
#         ).strip()

#     try:
#         data = json.loads(cleaned)
#         logger.info(f"Resume extracted for: {data.get('name', 'Unknown')}")
#         return data
#     except json.JSONDecodeError as e:
#         logger.error(f"JSON parse error after AI extraction: {e}\nRaw output: {content[:400]}")
#         return _fallback_data()


# def _fallback_data() -> dict:
#     return {
#         "name": "Candidate",
#         "email": None,
#         "phone": None,
#         "skills": [],
#         "experience": [],
#         "projects": [],
#         "education": [],
#         "summary": "Resume parsing encountered an issue. Proceeding with general interview questions.",
#     }


# # ── Main Entry Point ──────────────────────────────────────────────────────────

# def parse_resume(file_path: str) -> dict:
#     """
#     Full pipeline: file → text → Ollama → structured dict.

#     Returns:
#         Dict with keys: name, email, phone, skills, experience,
#                         projects, education, summary, raw_text
#     """
#     logger.info(f"Parsing resume: {file_path}")

#     raw_text = _extract_text(file_path)

#     if not raw_text.strip():
#         raise ValueError("Resume appears to be empty or could not be read.")

#     logger.info(f"Extracted {len(raw_text)} characters")

#     resume_data = _extract_with_ai(raw_text)
#     resume_data["raw_text"] = raw_text[:2000]  # Keep a snippet for reference
#     return resume_data

"""
resume/parser.py
================
Local resume parser without OpenAI.

Features:
- PDF extraction
- DOCX extraction
- Skill extraction
- Email extraction
- Phone extraction
- Basic resume structuring
"""

import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────
# Common Skills Database
# ─────────────────────────────────────────────────────────────
COMMON_SKILLS = [
    "Python",
    "Java",
    "JavaScript",
    "TypeScript",
    "React",
    "Next.js",
    "Node.js",
    "FastAPI",
    "Django",
    "Flask",
    "SQL",
    "MongoDB",
    "PostgreSQL",
    "MySQL",
    "AWS",
    "Docker",
    "Kubernetes",
    "Git",
    "HTML",
    "CSS",
    "Tailwind",
    "Selenium",
    "OpenAI",
    "Machine Learning",
    "AI",
    "REST API",
    "GraphQL",
    "Redis",
    "Linux",
    "C++",
    "C",
    "Go",
    "Rust",
]


# ─────────────────────────────────────────────────────────────
# PDF Extraction
# ─────────────────────────────────────────────────────────────
def extract_text_from_pdf(file_path: str) -> str:

    try:
        import PyPDF2

        text = []

        with open(file_path, "rb") as f:

            reader = PyPDF2.PdfReader(f)

            for page in reader.pages:

                page_text = page.extract_text()

                if page_text:
                    text.append(page_text)

        return "\n".join(text)

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")

        raise ValueError(f"Failed to parse PDF: {e}")


# ─────────────────────────────────────────────────────────────
# DOCX Extraction
# ─────────────────────────────────────────────────────────────
def extract_text_from_docx(file_path: str) -> str:

    try:
        from docx import Document

        doc = Document(file_path)

        text_parts = []

        for para in doc.paragraphs:

            if para.text.strip():
                text_parts.append(para.text.strip())

        return "\n".join(text_parts)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")

        raise ValueError(f"Failed to parse DOCX: {e}")


# ─────────────────────────────────────────────────────────────
# Route Extraction
# ─────────────────────────────────────────────────────────────
def extract_resume_text(file_path: str) -> str:

    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)

    elif extension in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)

    else:
        raise ValueError("Unsupported file type")


# ─────────────────────────────────────────────────────────────
# Extract Email
# ─────────────────────────────────────────────────────────────
def extract_email(text: str):

    match = re.search(
        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
        text
    )

    return match.group(0) if match else None


# ─────────────────────────────────────────────────────────────
# Extract Phone
# ─────────────────────────────────────────────────────────────
def extract_phone(text: str):

    match = re.search(
        r"(\+?\d[\d\s\-\(\)]{8,}\d)",
        text
    )

    return match.group(0) if match else None


# ─────────────────────────────────────────────────────────────
# Extract Skills
# ─────────────────────────────────────────────────────────────
def extract_skills(text: str):

    found_skills = []

    lower_text = text.lower()

    for skill in COMMON_SKILLS:

        if skill.lower() in lower_text:
            found_skills.append(skill)

    return sorted(list(set(found_skills)))


# ─────────────────────────────────────────────────────────────
# Extract Name
# ─────────────────────────────────────────────────────────────
def extract_name(text: str):

    lines = text.split("\n")

    for line in lines[:10]:

        line = line.strip()

        if (
            len(line.split()) >= 2
            and len(line.split()) <= 4
            and "resume" not in line.lower()
            and "cv" not in line.lower()
        ):
            return line

    return "Candidate"


# ─────────────────────────────────────────────────────────────
# Extract Education
# ─────────────────────────────────────────────────────────────
def extract_education(text: str):

    education_keywords = [
        "B.Tech",
        "Bachelor",
        "Master",
        "M.Tech",
        "MBA",
        "BCA",
        "MCA",
        "University",
        "College"
    ]

    education = []

    lines = text.split("\n")

    for line in lines:

        for keyword in education_keywords:

            if keyword.lower() in line.lower():

                education.append(line.strip())

                break

    return education[:3]


# ─────────────────────────────────────────────────────────────
# Extract Projects
# ─────────────────────────────────────────────────────────────
def extract_projects(text: str):

    projects = []

    lines = text.split("\n")

    capture = False

    for line in lines:

        if "project" in line.lower():
            capture = True
            continue

        if capture and line.strip():

            projects.append(line.strip())

            if len(projects) >= 5:
                break

    return projects


# ─────────────────────────────────────────────────────────────
# Main Parser
# ─────────────────────────────────────────────────────────────
def parse_resume(file_path: str) -> dict:

    logger.info(f"Parsing resume: {file_path}")

    raw_text = extract_resume_text(file_path)

    if not raw_text.strip():
        raise ValueError("Resume appears empty")

    resume_data = {
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "skills": extract_skills(raw_text),
        "experience": [],
        "projects": extract_projects(raw_text),
        "education": extract_education(raw_text),
        "summary": "Resume parsed successfully using local parser.",
        "raw_text": raw_text[:2000]
    }

    logger.info(
        f"Resume parsed successfully for: {resume_data['name']}"
    )

    return resume_data

